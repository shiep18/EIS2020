import docx
import pandas as pd
from openpyxl import Workbook
file = r'test.docx'
doc = docx.Document(file)
count = 0
tables = []
wb = Workbook()
ws = wb.active

table_num = len(doc.tables)
#table_num = 5
print(table_num)
for index in range(0, table_num):
  table = []
  for row in doc.tables[index].rows:
    line = []
    for grid in row.cells:
        
      line.append(grid.text)
    table.append(line)
    ws.append(line)
  count = count + 1
  print("第", count, "个表格正在处理...剩余", table_num - count + 1, "个表格", "\n")
  tables.append(table)
 
wb.save("全表.xlsx")
print(tables)
print("表格处理完成...")

